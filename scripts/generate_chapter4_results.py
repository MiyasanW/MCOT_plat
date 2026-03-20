from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "PROJECT_chapter4_results.docx"


def set_a4_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)


def set_font(run, size=14, bold=False):
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(size)
    run.bold = bold


def add_center_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=18, bold=True)


def add_heading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        set_font(r, size=16 if level == 2 else 14, bold=True)


def add_body(doc, text, indent=0.75):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        set_font(r, size=14)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            set_font(r, size=14)


def add_table(doc, title, headers, rows):
    add_heading(doc, title, level=3)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r, size=13, bold=True)

    for row_data in rows:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
            for p in row[i].paragraphs:
                for r in p.runs:
                    set_font(r, size=12)


def add_figure_placeholder(doc, figure_no, caption, detail):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[ใส่รูปภาพที่ {figure_no}: {caption}]")
    set_font(r, size=13, bold=True)

    d = doc.add_paragraph(detail)
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.paragraph_format.space_after = Pt(6)
    for run in d.runs:
        set_font(run, size=12)


def main():
    doc = Document()
    set_a4_margins(doc)

    add_center_title(doc, "บทที่ 4")
    add_center_title(doc, "ผลการดำเนินงาน")
    doc.add_paragraph()

    add_body(
        doc,
        "บทนี้นำเสนอผลการดำเนินงานของโครงงานระบบจองอุปกรณ์ถ่ายทำและสตูดิโอออนไลน์ โดยเน้นผลลัพธ์จากการพัฒนา "
        "การทดสอบ การนำขึ้นใช้งานจริง และผลกระทบต่อกระบวนการทำงานของหน่วยงาน เพื่อสะท้อนความพร้อมของระบบในการใช้งานต่อเนื่อง"
    )

    add_heading(doc, "4.1 ผลการพัฒนาเชิงฟังก์ชัน", level=2)
    add_body(
        doc,
        "ระบบที่พัฒนาขึ้นรองรับกระบวนการทำงานหลักได้ครบตั้งแต่ต้นทางถึงปลายทาง ได้แก่ การค้นหาและเลือกทรัพยากร "
        "การตรวจสอบความพร้อมใช้งาน การยืนยันการจอง การคำนวณยอดชำระและมัดจำ การอัปโหลดสลิปชำระเงิน และการจัดการสถานะโดยเจ้าหน้าที่"
    )
    add_bullets(
        doc,
        [
            "ลูกค้าสามารถสร้างรายการจองจากตะกร้าได้หลายรายการในครั้งเดียว",
            "ระบบตรวจสอบช่วงเวลาซ้ำซ้อนอัตโนมัติก่อนยืนยันการจอง",
            "ระบบส่งใบเสนอราคา (PDF) และอีเมลแจ้งเตือนตามเหตุการณ์สำคัญ",
            "เจ้าหน้าที่อนุมัติและเปลี่ยนสถานะได้ผ่านหน้า Staff Dashboard",
            "การติดตามสถานะของลูกค้าทำได้จากหน้า My Bookings แบบต่อเนื่อง",
        ],
    )
    add_figure_placeholder(
        doc,
        "4.1",
        "หน้าจอแคตตาล็อกและหน้ารายละเอียดทรัพยากร",
        "แสดงการค้นหา/กรองหมวดหมู่ และปุ่มเพิ่มลงตะกร้า"
    )
    add_figure_placeholder(
        doc,
        "4.2",
        "หน้าตะกร้าและขั้นตอนยืนยันการจอง",
        "แสดงรายการที่เลือก ช่วงวันเช่า ยอดรวม และยอดมัดจำ"
    )
    add_figure_placeholder(
        doc,
        "4.3",
        "หน้า Staff Dashboard สำหรับตรวจสอบและอนุมัติ",
        "แสดงรายการรอตรวจสอบ สถานะ และปุ่มดำเนินการ"
    )

    add_heading(doc, "4.2 ผลการทดสอบระบบ", level=2)
    add_body(
        doc,
        "ผลการทดสอบยืนยันว่าระบบทำงานได้ตามข้อกำหนดที่วางไว้ ทั้งในระดับตรรกะธุรกิจและระดับการใช้งานจริง "
        "โดยครอบคลุมการทดสอบการยืนยันตัวตน การจอง การยกเลิก และการสร้างเอกสาร"
    )

    add_table(
        doc,
        "ตาราง 4.1 สรุปผลการทดสอบกรณีสำคัญ",
        ["รหัส", "กรณีทดสอบ", "ผลที่คาดหวัง", "ผลทดสอบ"],
        [
            ["TC-01", "สร้างการจองจากตะกร้าปกติ", "สร้าง booking สำเร็จ", "ผ่าน"],
            ["TC-02", "มี id สินค้าไม่ถูกต้องในตะกร้า", "ระบบปฏิเสธพร้อมข้อความผิดพลาด", "ผ่าน"],
            ["TC-03", "ช่วงเวลาจองชนกัน", "ระบบแจ้งไม่พร้อมใช้งาน", "ผ่าน"],
            ["TC-04", "ผู้ใช้ทั่วไปเรียก staff endpoint", "ระบบปฏิเสธสิทธิ์", "ผ่าน"],
            ["TC-05", "สร้างใบเสนอราคา PDF", "ได้ไฟล์เอกสารครบถ้วน", "ผ่าน"],
        ],
    )
    add_figure_placeholder(
        doc,
        "4.4",
        "ผลการรันคำสั่งตรวจระบบ",
        "แสดงผลคำสั่ง python3 manage.py check ว่าไม่พบข้อผิดพลาด"
    )
    add_figure_placeholder(
        doc,
        "4.5",
        "ภาพหรือสกรีนช็อตผลการทดสอบสำคัญ",
        "แสดงตัวอย่าง test cases ที่ผ่าน เช่น booking, permission, cancellation"
    )

    add_heading(doc, "4.3 ผลการนำระบบขึ้นใช้งานจริง", level=2)
    add_body(
        doc,
        "หลังการนำระบบขึ้นใช้งานจริงบน VPS ระบบสามารถให้บริการได้ต่อเนื่องในเส้นทางหลัก โดยผ่านการตรวจสอบบริการ "
        "ที่เกี่ยวข้องและการทดสอบหลัง deploy ตามแนวปฏิบัติที่กำหนด"
    )
    add_bullets(
        doc,
        [
            "บริการแอปพลิเคชันและเว็บเซิร์ฟเวอร์ทำงานปกติหลัง restart",
            "หน้าเว็บหลักและ flow การจองใช้งานได้จริงในสภาพแวดล้อม production",
            "อีเมลแจ้งเตือนและเอกสารใบเสนอราคาทำงานได้ตามเงื่อนไข",
            "ผล smoke test และ health check อยู่ในเกณฑ์ที่ยอมรับได้",
        ],
    )
    add_figure_placeholder(
        doc,
        "4.6",
        "หน้าเว็บระบบหลังขึ้น production",
        "แสดงหน้าแรกหรือหน้าจองที่ทำงานจริงบนโดเมนระบบ"
    )
    add_figure_placeholder(
        doc,
        "4.7",
        "ผลตรวจสถานะบริการบน VPS",
        "แสดงสถานะ gunicorn/nginx หรือผล health check script"
    )

    add_heading(doc, "4.4 ผลลัพธ์เชิงประสิทธิภาพการทำงาน", level=2)
    add_body(
        doc,
        "เมื่อเปรียบเทียบกับกระบวนการเดิม ระบบช่วยให้การทำงานมีมาตรฐานมากขึ้น ลดงานซ้ำซ้อน และเพิ่มความชัดเจน "
        "ในการติดตามสถานะการจองระหว่างลูกค้าและเจ้าหน้าที่"
    )

    add_table(
        doc,
        "ตาราง 4.2 เปรียบเทียบผลก่อนและหลังพัฒนา",
        ["ประเด็น", "ก่อนพัฒนา", "หลังพัฒนา"],
        [
            ["การรับคำขอจอง", "หลายช่องทาง", "ช่องทางหลักเดียวผ่านระบบ"],
            ["การตรวจคิว/สต็อก", "ตรวจด้วยมือ", "ตรวจผ่านระบบอัตโนมัติ"],
            ["การติดตามสถานะ", "สอบถามซ้ำ", "ตรวจสอบได้จากระบบทันที"],
            ["การอนุมัติ", "ไม่เป็นมาตรฐาน", "มี workflow และสถานะชัดเจน"],
            ["การออกรายงาน", "รวบรวมข้อมูลล่าช้า", "เรียกดูข้อมูลได้รวดเร็วขึ้น"],
        ],
    )
    add_figure_placeholder(
        doc,
        "4.8",
        "กราฟหรือแดชบอร์ดสรุปจำนวนรายการจอง",
        "แสดงแนวโน้มจำนวนการจองและสถานะงาน เช่น Approved/Paid/Completed"
    )

    add_heading(doc, "4.5 ปัญหาที่พบระหว่างดำเนินงานและแนวทางแก้ไข", level=2)
    add_body(
        doc,
        "ระหว่างการพัฒนาพบประเด็นที่ต้องแก้ไขทั้งด้านข้อมูล ระบบ และขั้นตอนการทำงาน โดยสรุปปัญหาและแนวทางแก้ไขดังนี้"
    )

    add_bullets(
        doc,
        [
            "ปัญหาข้อมูลตะกร้าไม่สอดคล้องฐานข้อมูล: แก้โดยเพิ่ม validation ก่อนยืนยันทุกครั้ง",
            "ความแตกต่างระหว่าง local กับ production: แก้โดยมาตรฐานค่า environment และขั้นตอนตรวจหลัง deploy",
            "เงื่อนไขเปลี่ยนสถานะซับซ้อน: แก้โดยแยกตรรกะไป service layer และกำหนด workflow ชัดเจน",
            "รูปแบบการแจ้งเตือนยังไม่ชัดในบางกรณี: แก้โดยปรับข้อความและเหตุการณ์ที่ส่งอีเมล",
        ],
    )
    add_figure_placeholder(
        doc,
        "4.9",
        "ตัวอย่างอีเมลใบเสนอราคาและอีเมลแจ้งสถานะ",
        "แสดงหัวข้ออีเมล รายการที่แนบ และข้อความสำคัญที่ผู้ใช้ต้องทราบ"
    )

    add_heading(doc, "4.6 สรุปผลการดำเนินงานของบท", level=2)
    add_body(
        doc,
        "ผลการดำเนินงานในบทนี้แสดงให้เห็นว่าระบบสามารถตอบโจทย์การใช้งานจริงของหน่วยงานได้ ทั้งในด้านฟังก์ชันหลัก "
        "ความถูกต้องของกระบวนการ ความพร้อมใช้งานจริง และประสิทธิภาพการปฏิบัติงาน ระบบจึงมีความพร้อมสำหรับการใช้งานต่อเนื่อง "
        "และเป็นฐานสำหรับการพัฒนาขั้นถัดไป"
    )
    add_figure_placeholder(
        doc,
        "4.10",
        "ภาพรวม flow การทำงานตั้งแต่จองถึงปิดงาน",
        "สรุปลำดับขั้นตอน: ค้นหา -> จอง -> อนุมัติ -> ชำระ -> ส่งมอบ -> ปิดงาน"
    )

    add_heading(doc, "4.7 รายการภาพประกอบที่ต้องใส่", level=2)
    add_body(
        doc,
        "เพื่อให้การจัดรูปเล่มสมบูรณ์ ควรจัดเตรียมภาพประกอบตามรายการต่อไปนี้ โดยใช้หมายเลขภาพให้สอดคล้องกับเนื้อหาในบท"
    )
    add_bullets(
        doc,
        [
            "ภาพที่ 4.1: หน้าจอแคตตาล็อก/รายละเอียดทรัพยากร (วางในหัวข้อ 4.1)",
            "ภาพที่ 4.2: หน้าตะกร้าและยืนยันการจอง (วางในหัวข้อ 4.1)",
            "ภาพที่ 4.3: หน้า Staff Dashboard อนุมัติรายการ (วางในหัวข้อ 4.1)",
            "ภาพที่ 4.4: ผล python3 manage.py check (วางในหัวข้อ 4.2)",
            "ภาพที่ 4.5: ตัวอย่างผลชุดทดสอบสำคัญ (วางในหัวข้อ 4.2)",
            "ภาพที่ 4.6: หน้าเว็บ production หลัง deploy (วางในหัวข้อ 4.3)",
            "ภาพที่ 4.7: สถานะบริการบน VPS/health check (วางในหัวข้อ 4.3)",
            "ภาพที่ 4.8: กราฟหรือ dashboard สรุปผลลัพธ์ (วางในหัวข้อ 4.4)",
            "ภาพที่ 4.9: ตัวอย่างอีเมลแจ้งเตือน (วางในหัวข้อ 4.5)",
            "ภาพที่ 4.10: ภาพรวม flow การทำงานทั้งระบบ (วางท้ายหัวข้อ 4.6)",
        ],
    )

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
