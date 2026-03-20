from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "BIBLIOGRAPHY_APA7.docx"


def set_font(run, size=16, bold=False):
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(size)
    run.bold = bold


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run("บรรณานุกรม")
    set_font(rt, size=20, bold=True)

    refs = [
        "Django Software Foundation. (2026). Django documentation. https://docs.djangoproject.com/",
        "Python Software Foundation. (2026). Python documentation. https://docs.python.org/3/",
        "PostgreSQL Global Development Group. (2026). PostgreSQL documentation. https://www.postgresql.org/docs/",
        "Gunicorn Developers. (2026). Gunicorn documentation. https://docs.gunicorn.org/",
        "NGINX, Inc. (2026). NGINX documentation. https://nginx.org/en/docs/",
        "OWASP Foundation. (2026). OWASP top 10 web application security risks. https://owasp.org/www-project-top-ten/",
        "ReportLab. (2026). ReportLab user guide. https://www.reportlab.com/documentation/",
        "PyPDF2 Developers. (2026). PyPDF2 documentation. https://pypdf2.readthedocs.io/",
        "python-docx Developers. (2026). python-docx documentation. https://python-docx.readthedocs.io/",
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # APA 7 hanging indent
        p.paragraph_format.left_indent = Cm(0.0)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            set_font(run, size=16)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
