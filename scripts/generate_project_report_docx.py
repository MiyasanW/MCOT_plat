from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parents[1]
SOURCE_PATH = BASE_DIR / 'project_report.md'
OUTPUT_PATH = BASE_DIR / 'project_report.docx'
THAI_FONT = 'TH Sarabun New'
BODY_ALIGNMENT = getattr(WD_ALIGN_PARAGRAPH, 'THAI_JUSTIFY', WD_ALIGN_PARAGRAPH.JUSTIFY)
MAJOR_SECTIONS = {'บทคัดย่อ', 'บรรณานุกรม', 'ภาคผนวก'}


def set_document_defaults(document):
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.left_margin = Cm(3.81)

    normal_style = document.styles['Normal']
    normal_style.font.name = THAI_FONT
    normal_style.font.size = Pt(16)
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), THAI_FONT)

    settings = document.settings.element
    auto_hyphenation = settings.find(qn('w:autoHyphenation'))
    if auto_hyphenation is None:
        auto_hyphenation = OxmlElement('w:autoHyphenation')
        settings.append(auto_hyphenation)
    auto_hyphenation.set(qn('w:val'), '0')


def apply_run_font(run, size=16, bold=False):
    run.font.name = THAI_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), THAI_FONT)


def format_paragraph(paragraph, heading_level=0, list_item=False, code_block=False):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.left_indent = Cm(0)
    paragraph_format.right_indent = Cm(0)
    paragraph_format.keep_together = False
    paragraph_format.keep_with_next = False

    if heading_level > 0:
        paragraph_format.first_line_indent = Cm(0)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph_format.space_before = Pt(10)
        paragraph_format.space_after = Pt(6)
        paragraph_format.keep_with_next = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if heading_level <= 2 else WD_ALIGN_PARAGRAPH.LEFT
        return

    if list_item:
        paragraph_format.left_indent = Cm(0.75)
        paragraph_format.first_line_indent = Cm(-0.5)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return

    if code_block:
        paragraph_format.first_line_indent = Cm(0)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return

    paragraph_format.first_line_indent = Cm(1.25)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.alignment = BODY_ALIGNMENT


def add_heading(document, text, level, page_break_before=False):
    if page_break_before and document.paragraphs:
        document.add_page_break()

    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    size = {1: 24, 2: 20, 3: 18, 4: 16}.get(level, 16)
    apply_run_font(run, size=size, bold=True)
    format_paragraph(paragraph, heading_level=min(level, 3))

    if level == 4:
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_plain_paragraph(document, text, list_item=False, code_block=False):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    apply_run_font(run)
    format_paragraph(paragraph, list_item=list_item, code_block=code_block)
    if code_block:
        run.font.name = 'Courier New'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')


def add_numbered_item(document, text):
    dot_index = text.find('. ')
    paragraph = document.add_paragraph()
    number_run = paragraph.add_run(text[:dot_index + 1] + ' ')
    apply_run_font(number_run)
    content_run = paragraph.add_run(text[dot_index + 2:].strip())
    apply_run_font(content_run)
    format_paragraph(paragraph, list_item=True)


def add_table(document, table_lines):
    rows = []
    for line in table_lines:
        stripped = line.strip()
        if stripped.startswith('|') and stripped.endswith('|'):
            rows.append([cell.strip() for cell in stripped.strip('|').split('|')])

    if len(rows) < 2:
        for line in table_lines:
            add_plain_paragraph(document, line)
        return

    separator = rows[1]
    if all(not cell.replace(':', '').replace('-', '').strip() for cell in separator):
        rows.pop(1)

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = 'Table Grid'

    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                format_paragraph(paragraph)
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else BODY_ALIGNMENT
                for run in paragraph.runs:
                    apply_run_font(run, bold=row_index == 0)


def is_major_section_heading(text):
    return text.startswith('บทที่ ') or text in MAJOR_SECTIONS


def build_document():
    lines = SOURCE_PATH.read_text(encoding='utf-8').splitlines()
    document = Document()
    set_document_defaults(document)

    table_buffer = []
    in_table = False
    in_code_block = False

    for raw_line in lines:
        stripped = raw_line.strip()

        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            add_plain_paragraph(document, raw_line, code_block=True)
            continue

        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            table_buffer.append(raw_line)
            continue

        if in_table:
            add_table(document, table_buffer)
            table_buffer = []
            in_table = False

        if not stripped or stripped == '---':
            continue

        if raw_line.startswith('# '):
            add_heading(document, raw_line[2:].strip(), 1)
        elif raw_line.startswith('## '):
            heading_text = raw_line[3:].strip()
            add_heading(document, heading_text, 2 if is_major_section_heading(heading_text) else 3, page_break_before=is_major_section_heading(heading_text))
        elif raw_line.startswith('### '):
            add_heading(document, raw_line[4:].strip(), 4)
        elif raw_line.startswith('#### '):
            add_heading(document, raw_line[5:].strip(), 4)
        elif stripped.startswith('- '):
            add_plain_paragraph(document, stripped[2:].strip(), list_item=True)
        elif len(stripped) > 3 and stripped[0].isdigit() and '. ' in stripped[:4]:
            add_numbered_item(document, stripped)
        else:
            add_plain_paragraph(document, raw_line)

    if in_table and table_buffer:
        add_table(document, table_buffer)

    document.save(OUTPUT_PATH)
    print(f'CREATED {OUTPUT_PATH}')


if __name__ == '__main__':
    build_document()