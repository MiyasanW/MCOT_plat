from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "PROJECT_chapter2_introduction.docx"


def set_a4_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'TH Sarabun New'
        if level == 1:
            run.font.size = Pt(16)


def add_paragraph(doc, text, indent=0.75, is_sub=False):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if not is_sub:
        p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(14)


def main():
    doc = Document()
    set_a4_margins(doc)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("บทที่ 2")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = 'TH Sarabun New'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("บทนำ")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = 'TH Sarabun New'

    doc.add_paragraph()

    # 2.1 ที่มาและความสำคัญ
    add_heading(doc, "2.1 ที่มาและความสำคัญของปัญหา", level=2)

    add_paragraph(
        doc,
        "บริษัท อสมท จำกัด (มหาชน) ในฐานะองค์กรสื่อสารมวลชนของรัฐ มีการบริหารจัดการอุปกรณ์ถ่ายทำ "
        "สตูดิโอ และบริการที่เกี่ยวข้องกับงานผลิตสื่อจำนวนมากและหลากหลาย การจัดการทรัพยากรเหล่านี้ "
        "เบื้องต้นใช้วิธีการแบบแมนนวล อาทิ การขอคิวผ่านแชต การตรวจสอบสต็อกผ่านไฟล์ Excel และการติดตาม "
        "สถานะผ่านการโทรศัพท์หรือการพบปะตัวต่อตัว"
    )

    add_paragraph(
        doc,
        "ความท้าทายหลักของวิธีการเดิมประกอบด้วย (1) ข้อมูลไม่รวมศูนย์ ทำให้เกิดความขัดแย้งระหว่าง "
        "แหล่งข้อมูลหลายแห่ง (2) การจองซ้ำซ้อน เนื่องจากไม่มีการลัก off ข้อมูลแบบ real-time (3) ขั้นตอนการอนุมัติไม่มี "
        "มาตรฐาน ทำให้ติดต่อและสื่อสารล่าช้า (4) การตรวจสอบย้อนหลังและจัดทำรายงานทำได้ยากหากต้องการข้อมูลเชื่อมโยง "
        "ทั้งหลาย อีกทั้ง (5) ในช่วงที่หลายคนขอใช้งานพร้อมกัน (peak time) การจัดการจะเกิดความสับสน"
    )

    add_paragraph(
        doc,
        "จากสภาพปัญหาดังกล่าว หน่วยงานจึงมีความต้องการระบบกลาง (centralized system) ที่สามารถบริหารจัดการ "
        "การจองอุปกรณ์และสตูดิโอได้แบบครบวงจร ตั้งแต่ความต้องการของลูกค้า การตรวจสอบความพร้อมใช้งาน (availability) "
        "การยืนยันการจอง การชำระเงิน การอนุมัติจากผู้บริหาร และการติดตามสถานะอย่างต่อเนื่อง ระบบดังกล่าวคาดว่าจะช่วย "
        "ลดงานซ้ำซ้อน เพิ่มความโปร่งใส และเพิ่มประสิทธิภาพในการบริหารทรัพยากรขององค์กร"
    )

    doc.add_paragraph()

    # 2.2 วัตถุประสงค์
    add_heading(doc, "2.2 วัตถุประสงค์ของโครงงาน", level=2)

    add_heading(doc, "2.2.1 วัตถุประสงค์หลัก", level=3)

    objectives = [
        "พัฒนาระบบจองอุปกรณ์ถ่ายทำ สตูดิโอ และบริการ ที่รองรับการใช้งานจริงในองค์กร โดยให้ลูกค้า "
        "สามารถค้นหา ตรวจสอบความพร้อมใช้งาน และจองแบบหลายรายการได้พร้อมกัน",
        
        "ลดความผิดพลาดจากการจองซ้ำ ข้อมูลไม่ตรงกัน และปัญหาความสับสน โดยให้ข้อมูลรวมศูนย์ในระบบเดียว",
        
        "สร้างลำดับขั้นตอนการทำงานที่ชัดเจนสำหรับลูกค้า เจ้าหน้าที่ และผู้บริหาร โดยกำหนดสิทธิ์การเข้าถึง "
        "ตามบทบาท (Role-Based Access Control: RBAC)",
        
        "เพิ่มความสามารถในการติดตามสถานะการจอง การชำระเงิน และการอนุมัติอย่างต่อเนื่องและโปร่งใส",
        
        "เตรียมระบบให้รองรับการใช้งานจริง (production environment) ด้วยการปฏิบัติตามมาตรฐานด้าน "
        "ความปลอดภัย มาตรฐาน และการดูแลรักษา"
    ]

    for obj in objectives:
        add_paragraph(doc, obj, indent=0, is_sub=True)
        p = doc.paragraphs[-1]
        p.style = 'List Bullet'

    doc.add_paragraph()

    add_heading(doc, "2.2.2 วัตถุประสงค์รอง", level=3)

    secondary = [
        "ปรับปรุงประสบการณ์ผู้ใช้ (User Experience) ให้ใช้งานได้ดีทั้งบนอุปกรณ์พกพา (mobile) "
        "และเครื่องคอมพิวเตอร์ส่วนบุคคล (desktop) โดยใช้หลักการ mobile-first design",
        
        "วางมาตรฐานด้านความมั่นคงปลอดภัยเว็บแอปพลิเคชัน เช่น HTTPS, CSRF protection, "
        "การจัดการรหัสผ่าน และการตรวจสอบอนุญาต (authentication/authorization)",
        
        "จัดทำเอกสารครบถ้วน รวมถึง คู่มือผู้ใช้ (User Manual) เอกสารเทคนิค (Technical Documentation) "
        "และรายการตรวจสอบก่อนปล่อยระบบ (Release Checklist) สำหรับการถ่ายทอดและดูแลรักษา"
    ]

    for obj in secondary:
        add_paragraph(doc, obj, indent=0, is_sub=True)
        p = doc.paragraphs[-1]
        p.style = 'List Bullet'

    doc.add_paragraph()

    # 2.3 ขอบเขตของโครงงาน
    add_heading(doc, "2.3 ขอบเขตของโครงงาน", level=2)

    add_heading(doc, "2.3.1 ขอบเขตเชิงฟังก์ชัน (Functional Scope)", level=3)

    add_paragraph(doc, "รวมได้แก่:", indent=0, is_sub=True)

    functional = [
        "ระบบสมาชิก การลงทะเบียน การเข้าสู่ระบบ และการจัดการบัญชี",
        
        "แคตตาล็อกสินค้า สำหรับแสดงอุปกรณ์ สตูดิโอ แพ็กเกจ และบริการ พร้อมราคาและข้อมูลเกี่ยวข้อง",
        
        "ระบบตะกร้า (Shopping Cart) ที่อนุญาตให้ลูกค้าเลือกหลายรายการและทำการแก้ไขก่อนยืนยัน",
        
        "ระบบตรวจสอบความพร้อมใช้งาน (Availability Check) ตามช่วงวันที่เลือก เพื่อป้องกันการจองชนกัน",
        
        "ระบบสร้างรายการจอง (Booking Creation) พร้อมคำนวณราคา มัดจำ และ taxes/fees อัตโนมัติ",
        
        "ระบบแนบหลักฐานการชำระเงิน (Payment Slip Upload) พร้อมการตรวจสอบและยืนยัน",
        
        "ระบบฝั่งเจ้าหน้าที่ (Staff Dashboard) สำหรับจัดการ อนุมัติ และติดตามสถานะรายการจอง",
        
        "ระบบสร้างเอกสาร PDF สำหรับใบเสนอราคา (Quotation) ใบปิดบิล (Invoice) และเอกสารชำระเงินอื่น ๆ"
    ]

    for item in functional:
        add_paragraph(doc, item, indent=0, is_sub=True)
        p = doc.paragraphs[-1]
        p.style = 'List Bullet'

    doc.add_paragraph()

    add_heading(doc, "2.3.2 ขอบเขตเชิงเทคนิค (Technical Scope)", level=3)

    add_paragraph(doc, "รวมได้แก่:", indent=0, is_sub=True)

    technical = [
        "พัฒนาเป็นเว็บแอปพลิเคชัน (Web Application) โดยใช้ Django Framework (Python) "
        "สำหรับฝั่งหลังบ้าน (Backend) และ HTML/CSS/JavaScript สำหรับฝั่งหน้าบ้าน (Frontend)",
        
        "รองรับการนำระบบขึ้นใช้งานจริง (Deployment) บนเครื่องแม่ข่ายเสมือน (VPS) "
        "ด้วย Gunicorn และ Nginx",
        
        "รองรับการตั้งค่าต่างระหว่างสภาพแวดล้อมพัฒนา (Development) และ production "
        "ผ่าน environment variables",
        
        "มีชุดคำสั่งตรวจสุขภาพระบบ (Health Check) และทดสอบหลังนำขึ้นใช้งานจริง (Smoke Test)"
    ]

    for item in technical:
        add_paragraph(doc, item, indent=0, is_sub=True)
        p = doc.paragraphs[-1]
        p.style = 'List Bullet'

    doc.add_paragraph()

    add_heading(doc, "2.3.3 ขอบเขตที่ไม่รวมในโครงงาน", level=3)

    add_paragraph(doc, "ไม่รวมได้แก่:", indent=0, is_sub=True)

    excluded = [
        "ระบบชำระเงินออนไลน์แบบ Payment Gateway เต็มรูปแบบ (เช่น Stripe หรือ 2C2P) "
        "ในขั้นตอนนี้ยังเป็นการอัปโหลดสลิปแบบแมนนวล)",
        
        "Mobile Application แบบ Native (iOS/Android) - สั่งให้ใช้เว็บแบบ responsive แทน",
        
        "ระบบ Business Intelligence (BI) เชิงลึกหรือ Data Warehouse ระดับองค์กร - "
        "มีเพียง reporting พื้นฐานเท่านั้น",
        
        "ระบบเชื่อมต่อกับการลงทะเบียนหรือระบบบัญชีเงินกลางอื่นขององค์กร "
        "(integration เพื่อให้สอดคล้องกับทีมอื่น แต่ไม่เป็นส่วนแก่นของโครงงาน)"
    ]

    for item in excluded:
        add_paragraph(doc, item, indent=0, is_sub=True)
        p = doc.paragraphs[-1]
        p.style = 'List Bullet'

    doc.add_paragraph()

    # 2.4 ประโยชน์ที่คาดว่าจะได้รับ
    add_heading(doc, "2.4 ประโยชน์ที่คาดว่าจะได้รับ", level=2)

    add_heading(doc, "2.4.1 ประโยชน์ต่อหน่วยงาน", level=3)

    benefits_org = [
        "ลดเวลาการประสานงาน เจ้าหน้าที่ไม่ต้องตอบคำติดต่อหลายรูปแบบซ้ำ ๆ",
        
        "เพิ่มความถูกต้องของข้อมูลการจอง ลด human error และความสับสนจากหลายช่องทาง",
        
        "มีระบบติดตามสถานะแบบรวมศูนย์ สามารถออกรายงานทันที และตรวจสอบย้อนหลังได้",
        
        "วางมาตรฐานการปฏิบัติงาน ทำให้การส่งมอบบริการเสมอภาค"
    ]

    for item in benefits_org:
        add_paragraph(doc, item, indent=0, is_sub=True)
        p = doc.paragraphs[-1]
        p.style = 'List Bullet'

    doc.add_paragraph()

    add_heading(doc, "2.4.2 ประโยชน์ต่อผู้ใช้งาน (ลูกค้า)", level=3)

    benefits_user = [
        "จองทรัพยากรได้สะดวก ไม่ต้องติดต่อเจ้าหน้าที่หลายครั้ง",
        
        "ตรวจสอบสถานะรายการของตนเองได้ต่อเนื่อง ไม่ต้องรอการสอบถาม",
        
        "ลดความไม่แน่นอนจากการสื่อสารหลายช่องทาง ข้อมูลเป็นหนึ่งเดียว"
    ]

    for item in benefits_user:
        add_paragraph(doc, item, indent=0, is_sub=True)
        p = doc.paragraphs[-1]
        p.style = 'List Bullet'

    doc.add_paragraph()

    add_heading(doc, "2.4.3 ประโยชน์ต่อผู้จัดทำสหกิจศึกษา", level=3)

    benefits_intern = [
        "ได้ฝึกพัฒนาระบบทั้งฝั่งหน้าบ้านและหลังบ้าน (full-stack) ในงานจริง",
        
        "ได้เรียนรู้การนำระบบขึ้นใช้งานจริง (production deployment) และการดูแลในสภาพแวดล้อมจริง",
        
        "ได้ประสบการณ์การทำงานร่วมกับผู้ใช้งานจริงและการบริหารจัดการความต้องการที่เปลี่ยนแปลง",
        
        "ได้จัดทำเอกสารและคู่มือที่ใช้ประจำ (production-ready documentation)"
    ]

    for item in benefits_intern:
        add_paragraph(doc, item, indent=0, is_sub=True)
        p = doc.paragraphs[-1]
        p.style = 'List Bullet'

    doc.add_paragraph()

    # 2.5 นิยามศัพท์เฉพาะ
    add_heading(doc, "2.5 นิยามศัพท์เฉพาะ", level=2)

    add_paragraph(
        doc,
        "ต่อไปนี้คือนิยามของศัพท์ที่ใช้ในเอกสารโครงงานนี้ เพื่อให้ความเข้าใจเป็นไปในทิศทางเดียวกัน"
    )

    doc.add_paragraph()

    # Terminology Table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'

    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = "ศัพท์"
    header_cells[1].text = "นิยาม"

    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'TH Sarabun New'

    terms = [
        ["Availability", "ความพร้อมใช้งาน หรือช่วงเวลาที่ทรัพยากรว่างสำหรับจอง"],
        ["Booking", "การจอง หรือรายการคำขอใช้ทรัพยากรจากลูกค้า"],
        ["Dashboard", "หน้าแสดงข้อมูลสรุป หรือแผงควบคุม สำหรับรีงบริหาร"],
        ["Deployment", "การนำระบบขึ้นใช้งานจริง หรือปล่อยให้ใช้ในสภาพแวดล้อมจริง"],
        ["Permission", "สิทธิ์การเข้าถึง หรือความสามารถในการใช้ฟังก์ชันบางอย่าง"],
        ["Role-Based Access Control (RBAC)", "ระบบควบคุมสิทธิ์ตามบทบาท เช่น admin, staff, customer"],
        ["Status", "สถานะ หรือสภาพปัจจุบันของรายการจอง เช่น pending, approved, paid"],
        ["Deposit", "มัดจำ หรือจำนวนเงินที่ลูกค้าต้องชำระก่อนนำทรัพยากรไป"],
        ["Quotation", "ใบเสนอราคา เอกสารแสดงรายละเอียดราคาและเงื่อนไขของการจอง"],
        ["Notification", "การแจ้งเตือน เช่น อีเมลหรือข้อความแจ้งสถานะการจอง"],
    ]

    for term, definition in terms:
        row = table.add_row()
        row.cells[0].text = term
        row.cells[1].text = definition

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'TH Sarabun New'
                    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f"✅ สร้าง {OUTPUT} สำเร็จ")


if __name__ == "__main__":
    main()
