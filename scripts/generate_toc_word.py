from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER

OUT = "TOC_from_Thanandorn_DSM16541N.docx"


def set_font(run, size=16, bold=False):
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(size)
    run.bold = bold


def add_toc_line(doc, text, page, indent_cm=0.0, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(16.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )

    r1 = p.add_run(text)
    set_font(r1, size=16, bold=bold)

    page_text = page if page else "-"
    r2 = p.add_run("\t" + page_text)
    set_font(r2, size=16, bold=bold)


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.1)
    sec.bottom_margin = Cm(2.1)
    sec.left_margin = Cm(2.1)
    sec.right_margin = Cm(2.1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run("สารบัญ")
    set_font(rt, size=20, bold=True)

    entries = [
        ("บทคัดย่อภาษาไทย", "ก", 0.0, False),
        ("บทคัดย่อภาษาอังกฤษ", "ข", 0.0, False),
        ("กิตติกรรมประกาศ", "ค", 0.0, False),
        ("สารบัญ", "ง", 0.0, False),
        ("สารบัญตาราง", "จ", 0.0, False),
        ("สารบัญภาพ", "ฉ", 0.0, False),
        ("บทที่ 1 ข้อมูลเกี่ยวกับสถานประกอบการ", "1", 0.0, True),
        ("1.1 ชื่อ ที่ตั้งและลักษณะการดําเนินงาน", "1", 0.7, False),
        ("1.2 วิสัยทัศน์ พันธกิจและค่านิยมของบริษัท อสมท จํากัด (มหาชน)", "2", 0.7, False),
        ("1.3 ค่านิยมของบริษัท อสมท จํากัด (มหาชน)", "2", 0.7, False),
        ("1.4 โครงสร้างของบริษัท อสมท จํากัด (มหาชน)", "3", 0.7, False),
        ("1.5 ระยะเวลาปฏิบัติงาน", "3", 0.7, False),
        ("บทที่ 2 บทนํา", "4", 0.0, True),
        ("2.1 ที่มาและความสําคัญของปัญหา", "4", 0.7, False),
        ("2.2 วัตถุประสงค์", "5", 0.7, False),
        ("2.3 ประโยชน์ที่คาดว่าจะได้รับ", "5", 0.7, False),
        ("2.4 ขอบเขตการทําโครงงาน", "5", 0.7, False),
        ("2.5 นิยามศัพท์เฉพาะ", "5", 0.7, False),
        ("บทที่ 3 ทฤษฎีและเทคโนโลยีที่ใช้ในการปฏิบัติงาน", "13", 0.0, True),
        ("3.1 แนวคิดและทฤษฎีที่เกี่ยวข้อง", "13", 0.7, False),
        ("3.2 เครื่องมือที่ใช้ในการพัฒนา", "14", 0.7, False),
        ("3.3 การวิเคราะห์กระบวนการทำงาน", "15", 0.7, False),
        ("3.4 การออกแบบสถาปัตยกรรมระบบ", "16", 0.7, False),
        ("3.5 การออกแบบฐานข้อมูล", "17", 0.7, False),
        ("3.6 เครื่องมือและภาษาที่ใช้ในการพัฒนา", "17", 0.7, False),
        ("3.7 การพัฒนาและการทดสอบ", "18", 0.7, False),
        ("3.8 การนำระบบขึ้นใช้งานจริง", "18", 0.7, False),
        ("3.9 สรุปวิธีการดำเนินงาน", "18", 0.7, False),

        ("บทที่ 4 ผลการปฏิบัติงาน", "19", 0.0, True),
        ("4.1 ผลการพัฒนาเชิงฟังก์ชัน", "19", 0.7, False),
        ("4.2 ผลการทดสอบระบบ", "23", 0.7, False),
        ("4.3 ผลการนำระบบขึ้นใช้งานจริง", "24", 0.7, False),
        ("4.4 ผลลัพธ์เชิงประสิทธิภาพการทำงาน", "24", 0.7, False),
        ("4.5 ปัญหาที่พบระหว่างดำเนินงานและแนวทางแก้ไข", "25", 0.7, False),
        ("4.6 สรุปผลการดำเนินงานของบท", "26", 0.7, False),
        ("4.7 รายการภาพประกอบที่ต้องใส่", "26", 0.7, False),

        ("บทที่ 5 สรุปผล อภิปรายผล และข้อเสนอแนะ", "27", 0.0, True),
        ("5.1 สรุปผล", "27", 0.7, False),
        ("5.2 อภิปรายผล", "27", 0.7, False),
        ("5.3 วิเคราะห์ปัญหา สาเหตุ และแนวทางแก้ไข", "28", 0.7, False),
        ("5.3.1 วิเคราะห์ปัญหาที่เกิดขึ้น", "28", 1.3, False),
        ("5.3.2 สาเหตุของปัญหา", "28", 1.3, False),
        ("5.3.3 แนวทางการแก้ไข", "29", 1.3, False),
        ("5.4 ข้อเสนอแนะ", "29", 0.7, False),
        ("บรรณานุกรม", "30", 0.0, False),
        ("ภาคผนวก ก ใบบันทึกลงเวลาปฏิบัติงาน", "31", 0.0, False),
        ("ภาคผนวก ข แบบบันทึกการปฏิบัติงานสหกิจศึกษา", "32", 0.0, False),
        ("ประวัติผู้จัดทํา", "-", 0.0, False),
    ]

    for text, page, indent, bold in entries:
        add_toc_line(doc, text, page, indent_cm=indent, bold=bold)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
