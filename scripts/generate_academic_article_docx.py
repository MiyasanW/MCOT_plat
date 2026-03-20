from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "academic_article_mcot_no_abstract.docx"


def set_a4_and_margins(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)


def set_columns(section, num=2, space_twips=340):
    sect_pr = section._sectPr
    cols = sect_pr.xpath('./w:cols')
    if cols:
        col = cols[0]
    else:
        col = OxmlElement('w:cols')
        sect_pr.append(col)

    col.set(qn('w:num'), str(num))
    col.set(qn('w:space'), str(space_twips))


def add_heading(doc, text, size=14):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(size)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)
    return p


def main():
    doc = Document()
    set_a4_and_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = "TH Sarabun New"
    style.font.size = Pt(14)

    # Header block (single column)
    p = doc.add_paragraph("การพัฒนาระบบบริหารการเช่าอุปกรณ์และสตูดิโอผ่านเว็บแอปพลิเคชันสำหรับงานผลิตสื่อ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(18)

    p = doc.add_paragraph(
        "กรณีศึกษา MCOT Equipment Service"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(16)
    p.runs[0].bold = True

    p = doc.add_paragraph(
        "Development of a Web-Based Equipment and Studio Rental Management System "
        "for Media Production: A Case Study of MCOT Equipment Service"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(13)

    p = doc.add_paragraph("ชื่อผู้เขียนคนที่ 1, ชื่อผู้เขียนคนที่ 2")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("สาขาวิชาเทคโนโลยีสารสนเทศ/วิศวกรรมซอฟต์แวร์ คณะ... มหาวิทยาลัย...")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("E-mail: ...")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Main article body in two columns (similar to sample)
    set_columns(doc.sections[0], num=2, space_twips=340)

    add_heading(doc, "บทคัดย่อ", size=16)
    add_body_paragraph(
        doc,
        "โครงงานนี้มีวัตถุประสงค์เพื่อพัฒนาระบบ MCOT Equipment Service สำหรับบริหารการเช่าอุปกรณ์ "
        "สตูดิโอ แพ็กเกจ และบริการที่เกี่ยวข้องกับงานผลิตสื่อในรูปแบบเว็บแอปพลิเคชัน โดยครอบคลุมกระบวนการ"
        "ตั้งแต่การค้นหาและเลือกทรัพยากร การตรวจสอบความพร้อมใช้งาน การสร้างรายการจอง การคำนวณราคา"
        "และมัดจำ การอัปโหลดหลักฐานการชำระเงิน การอนุมัติโดยเจ้าหน้าที่ และการติดตามสถานะรายการ"
    )
    add_body_paragraph(
        doc,
        "ผลการดำเนินงานพบว่าระบบสามารถใช้งานได้จริงตามขอบเขตที่กำหนด ช่วยลดงานซ้ำซ้อนจากการ"
        "ประสานงานหลายช่องทาง ลดความคลาดเคลื่อนของข้อมูล และเพิ่มความรวดเร็วในการติดตามสถานะ"
        "ทั้งยังรองรับการนำขึ้นใช้งานจริงและพร้อมต่อยอดในระดับองค์กร"
    )
    add_body_paragraph(
        doc,
        "คำสำคัญ: ระบบจองทรัพยากร, การเช่าอุปกรณ์, เว็บแอปพลิเคชัน, Django, การบริหารงานเช่า"
    )

    add_heading(doc, "1. คำนำ", size=15)
    add_body_paragraph(
        doc,
        "การบริหารการเช่าอุปกรณ์และสตูดิโอสำหรับงานผลิตสื่อมีความซับซ้อนจากผู้ใช้งานหลายบทบาทและ"
        "เงื่อนไขการจองที่หลากหลาย โครงงานนี้จึงพัฒนาระบบกลางเพื่อรวมกระบวนการทั้งหมดให้อยู่ในแพลตฟอร์มเดียว"
        "และตรวจสอบย้อนหลังได้"
    )

    add_heading(doc, "2. ทฤษฎีและเทคโนโลยีที่เกี่ยวข้อง", size=15)
    add_heading(doc, "2.1 แนวคิดการออกแบบระบบ", size=14)
    add_body_paragraph(
        doc,
        "ระบบใช้แนวคิด Role-Based Access Control และออกแบบ workflow การจองแบบลำดับขั้น "
        "โดยมีสถานะหลัก ได้แก่ Draft, Pending, Approved, Active, Overdue, Completed และ Cancelled"
    )
    add_heading(doc, "2.2 เทคโนโลยีที่ใช้", size=14)
    add_body_paragraph(
        doc,
        "ระบบพัฒนาด้วย Django Framework และฐานข้อมูลเชิงสัมพันธ์ พร้อมกลไกแจ้งเตือน การจัดการเอกสาร "
        "และมาตรการความมั่นคงปลอดภัยสำหรับการใช้งานจริง"
    )

    add_heading(doc, "3. ขั้นตอนและวิธีดำเนินงาน", size=15)
    add_heading(doc, "3.1 การวิเคราะห์ความต้องการ", size=14)
    add_body_paragraph(
        doc,
        "เริ่มจากการศึกษาปัญหากระบวนการเดิมและนิยามขอบเขตระบบ ก่อนออกแบบข้อมูลและโฟลว์ของผู้ใช้งาน"
    )
    add_heading(doc, "3.2 การพัฒนาและทดสอบ", size=14)
    add_body_paragraph(
        doc,
        "ดำเนินงานแบบ iterative เริ่มจากวิเคราะห์ความต้องการ ออกแบบข้อมูลและหน้าจอ พัฒนาฟังก์ชันหลัก "
        "ทดสอบเชิงฟังก์ชัน และทดสอบหลัง deploy"
    )

    add_heading(doc, "4. ผลการดำเนินงาน", size=15)
    add_body_paragraph(
        doc,
        "ระบบรองรับกระบวนการจองตั้งแต่ต้นจนจบ ลดงานซ้ำซ้อน เพิ่มความโปร่งใสในการติดตามสถานะ "
        "และช่วยให้เจ้าหน้าที่จัดการรายการได้เป็นมาตรฐาน"
    )

    add_heading(doc, "5. สรุป อภิปรายผล และข้อเสนอแนะ", size=15)
    add_body_paragraph(
        doc,
        "ระบบที่พัฒนาช่วยยกระดับการบริหารงานเช่าให้เป็นดิจิทัลอย่างมีประสิทธิภาพ และควรต่อยอด"
        "ด้านรายงานเชิงวิเคราะห์และการเชื่อมต่อระบบภายนอกในอนาคต"
    )

    add_heading(doc, "6. กิตติกรรมประกาศ", size=15)
    add_body_paragraph(
        doc,
        "ผู้จัดทำขอขอบคุณบริษัท อสมท จำกัด (มหาชน) และอาจารย์ที่ปรึกษาที่ให้คำแนะนำตลอดการดำเนินโครงงาน"
    )

    add_heading(doc, "เอกสารอ้างอิง", size=15)
    add_body_paragraph(doc, "[1] Django Software Foundation, Django Documentation.")
    add_body_paragraph(doc, "[2] OWASP Foundation, OWASP Top 10 Web Application Security Risks.")
    add_body_paragraph(doc, "[3] Python Software Foundation, Python Documentation.")
    add_body_paragraph(doc, "[4] เอกสารภายในโครงการ: DOCUMENTATION, USER_MANUAL, RELEASE_CHECKLIST, README.")

    add_heading(doc, "ประวัติผู้เขียนโดยสังเขป", size=15)
    add_body_paragraph(doc, "ชื่อ-สกุล: ...")
    add_body_paragraph(doc, "สาขาวิชา: ...")
    add_body_paragraph(doc, "คณะ/มหาวิทยาลัย: ...")
    add_body_paragraph(doc, "E-mail: ...")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
