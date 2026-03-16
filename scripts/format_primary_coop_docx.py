import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parents[1]
TARGET_PATH = BASE_DIR / 'โครงงานสหกิจศึกษาธนันดร_ฉบับรวม.docx'
BACKUP_PATH = BASE_DIR / 'โครงงานสหกิจศึกษาธนันดร_ฉบับรวม.backup.docx'
THAI_FONT = 'TH Sarabun New'
BODY_ALIGNMENT = getattr(WD_ALIGN_PARAGRAPH, 'THAI_JUSTIFY', WD_ALIGN_PARAGRAPH.JUSTIFY)

MAJOR_EXACT_HEADINGS = {
    'บทคัดย่อ',
    'Abstract',
    'กิตติกรรมประกาศ',
    'สารบัญ',
    'บรรณานุกรม',
    'ภาคผนวก',
}
CHAPTER_PATTERN = re.compile(r'^บทที่\s*\d+(?:\s+.*)?$')
STANDALONE_CHAPTER_PATTERN = re.compile(r'^บทที่\s*\d+$')
SUBHEADING_PATTERN = re.compile(r'^\d+(?:\.\d+)+\s+')
NUMBERED_ITEM_PATTERN = re.compile(r'^\d+\.\s+')
CONTINUATION_START_PATTERN = re.compile(r'^(และ|โดย|ซึ่ง|นอกจากนี้|ดังนั้น|ทั้งนี้|รวมถึง|ได้|ที่|ใน|จาก|เพื่อ)')
TERMINAL_END_PATTERN = re.compile(r'[\.!\?…:;"”’\)]$')


def set_document_defaults(document):
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.left_margin = Cm(3.81)

    normal_style = document.styles['Normal']
    normal_style.font.name = THAI_FONT
    normal_style.font.size = Pt(16)
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), THAI_FONT)

    settings = document.settings.element
    auto_hyphenation = settings.find(qn('w:autoHyphenation'))
    if auto_hyphenation is None:
        auto_hyphenation = OxmlElement('w:autoHyphenation')
        settings.append(auto_hyphenation)
    auto_hyphenation.set(qn('w:val'), '0')


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    apply_run_font_placeholder(run)

    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '

    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def apply_run_font_placeholder(run):
    run.font.name = THAI_FONT
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), THAI_FONT)


def apply_run_font(paragraph, size=16, bold=None):
    for run in paragraph.runs:
        run.font.name = THAI_FONT
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), THAI_FONT)
        if bold is not None:
            run.bold = bold


def deduplicate_chapter_blocks(document):
    paragraphs = list(document.paragraphs)
    indices_to_remove = set()

    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if not STANDALONE_CHAPTER_PATTERN.match(text):
            continue

        title_index = index + 1
        if title_index >= len(paragraphs):
            continue
        title_text = paragraphs[title_index].text.strip()
        if not title_text or title_text == '---':
            continue

        combined_heading = f'{text} {title_text}'.strip()
        match_index = None
        for look_ahead in range(title_index + 1, min(index + 15, len(paragraphs))):
            if paragraphs[look_ahead].text.strip() == combined_heading:
                match_index = look_ahead
                break

        if match_index is None:
            continue

        for remove_index in range(index, match_index):
            indices_to_remove.add(remove_index)

    for index in sorted(indices_to_remove, reverse=True):
        remove_paragraph(paragraphs[index])


def remove_separator_paragraphs(document):
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == '---']
    for paragraph in paragraphs:
        remove_paragraph(paragraph)


def ensure_footer_page_numbers(document):
    for section in document.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        for paragraph in list(footer.paragraphs):
            if paragraph.text.strip():
                remove_paragraph(paragraph)
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        add_page_number(paragraph)


def is_major_heading(text):
    return text in MAJOR_EXACT_HEADINGS or bool(CHAPTER_PATTERN.match(text))


def is_subheading(text):
    return bool(SUBHEADING_PATTERN.match(text))


def is_numbered_item(text):
    return bool(NUMBERED_ITEM_PATTERN.match(text))


def is_caption_or_label(text):
    return text.startswith(('ภาพที่', 'ตาราง', 'หมายเหตุ:', 'ชื่อโครงงาน', 'Project Name', 'By ', 'Workplace', 'Adviser', 'Major Fied', 'Academic Year'))


def looks_like_body_paragraph(text):
    if len(text) < 40:
        return False
    if is_caption_or_label(text):
        return False
    if is_major_heading(text) or is_subheading(text) or is_numbered_item(text):
        return False
    if text.endswith(':'):
        return False
    return True


def should_merge_with_previous(previous_text, current_text):
    if not previous_text or not current_text:
        return False
    if is_major_heading(previous_text) or is_major_heading(current_text):
        return False
    if is_subheading(previous_text) or is_subheading(current_text):
        return False
    if is_numbered_item(previous_text) or is_numbered_item(current_text):
        return False
    if is_caption_or_label(previous_text) or is_caption_or_label(current_text):
        return False
    if current_text.endswith(':'):
        return False

    # Merge line-break artifacts where the next paragraph is clearly a continuation.
    if len(current_text) <= 12:
        return True
    if CONTINUATION_START_PATTERN.match(current_text):
        return True

    # If previous paragraph ends with an incomplete short token, treat next line as continuation.
    prev_last_word = previous_text.split()[-1] if previous_text.split() else ''
    if prev_last_word and len(prev_last_word) <= 4 and not TERMINAL_END_PATTERN.search(previous_text):
        return True

    return False


def merge_fragmented_paragraphs(document):
    changed = True
    while changed:
        changed = False
        paragraphs = list(document.paragraphs)
        for index in range(1, len(paragraphs)):
            previous = paragraphs[index - 1]
            current = paragraphs[index]
            previous_text = previous.text.strip()
            current_text = current.text.strip()
            if not previous_text or not current_text:
                continue
            if should_merge_with_previous(previous_text, current_text):
                joiner = '' if previous_text.endswith(' ') else ' '
                previous.text = previous_text + joiner + current_text
                remove_paragraph(current)
                changed = True
                break


def format_paragraph(paragraph, text, previous_text):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Cm(0)
    paragraph_format.right_indent = Cm(0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.keep_together = False
    paragraph_format.keep_with_next = False
    paragraph_format.page_break_before = False

    if is_major_heading(text):
        paragraph_format.page_break_before = True
        paragraph_format.first_line_indent = Cm(0)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph_format.space_before = Pt(12)
        paragraph_format.space_after = Pt(8)
        paragraph_format.keep_with_next = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_run_font(paragraph, size=20, bold=True)
        return

    if previous_text and CHAPTER_PATTERN.match(previous_text) and text and not is_subheading(text) and not is_numbered_item(text) and len(text) < 80:
        paragraph_format.first_line_indent = Cm(0)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(8)
        paragraph_format.keep_with_next = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_run_font(paragraph, size=18, bold=True)
        return

    if is_subheading(text):
        paragraph_format.first_line_indent = Cm(0)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(2)
        paragraph_format.keep_with_next = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        apply_run_font(paragraph, size=16, bold=True)
        return

    if paragraph.style.name == 'List Paragraph' or is_numbered_item(text):
        paragraph_format.left_indent = Cm(0.75)
        paragraph_format.first_line_indent = Cm(-0.5)
        paragraph_format.line_spacing = 1.5
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        apply_run_font(paragraph, size=16)
        return

    if looks_like_body_paragraph(text):
        paragraph_format.first_line_indent = Cm(1.25)
        paragraph_format.line_spacing = 1.5
        paragraph.alignment = BODY_ALIGNMENT
        paragraph_format.widow_control = True
        apply_run_font(paragraph, size=16)
        return

    paragraph_format.first_line_indent = Cm(0)
    paragraph_format.line_spacing = 1.5
    paragraph.alignment = paragraph.alignment or WD_ALIGN_PARAGRAPH.LEFT
    apply_run_font(paragraph, size=16)


def main():
    if not TARGET_PATH.exists():
        raise FileNotFoundError(f'Missing target file: {TARGET_PATH}')

    if not BACKUP_PATH.exists():
        shutil.copy2(TARGET_PATH, BACKUP_PATH)

    document = Document(TARGET_PATH)
    set_document_defaults(document)
    deduplicate_chapter_blocks(document)
    remove_separator_paragraphs(document)
    merge_fragmented_paragraphs(document)
    ensure_footer_page_numbers(document)

    previous_text = ''
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text or text == '---':
            previous_text = text
            continue
        format_paragraph(paragraph, text, previous_text)
        previous_text = text

    document.save(TARGET_PATH)
    print(f'FORMATTED {TARGET_PATH}')
    print(f'BACKUP {BACKUP_PATH}')


if __name__ == '__main__':
    main()