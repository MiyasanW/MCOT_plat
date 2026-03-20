from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "APPENDIX_USER_MANUAL.docx"


def set_page(doc):
    s = doc.sections[0]
    s.top_margin = Cm(2.1)
    s.bottom_margin = Cm(2.1)
    s.left_margin = Cm(2.1)
    s.right_margin = Cm(2.1)


def style_run(run, size=14, bold=False):
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(size)
    run.bold = bold


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    style_run(r, size=18, bold=True)


def add_heading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        style_run(r, size=16 if level == 2 else 14, bold=True)


def add_body(doc, text, indent=0.75):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        style_run(r, size=14)


def add_list(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for r in p.runs:
            style_run(r, size=14)


def add_code(doc, code):
    p = doc.add_paragraph(code)
    p.paragraph_format.left_indent = Cm(0.75)
    for r in p.runs:
        style_run(r, size=12)


def add_image_placeholder(doc, no, caption, where):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[ใส่ภาพ ช.{no}: {caption}]")
    style_run(r, size=13, bold=True)

    p2 = doc.add_paragraph(f"ตำแหน่งแนะนำ: {where}")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p2.runs:
        style_run(r, size=12)


def main():
    doc = Document()
    set_page(doc)

    add_title(doc, "ภาคผนวก ช")
    add_title(doc, "คู่มือการใช้งานระบบ MCOT Equipment Service")

    add_body(doc, "เอกสารฉบับนี้จัดทำขึ้นเพื่อใช้เป็นคู่มือประกอบการใช้งานระบบสำหรับแนบในเล่มโครงงาน ครอบคลุมขั้นตอนปฏิบัติจริงสำหรับผู้ใช้งานทุกบทบาท ได้แก่ ลูกค้า เจ้าหน้าที่ และผู้ดูแลระบบ เพื่อให้การใช้งานเป็นมาตรฐานเดียวกันและตรวจสอบได้")

    add_heading(doc, "ช.1 วัตถุประสงค์ของคู่มือ", level=2)
    add_list(doc, [
        "อธิบายขั้นตอนการใช้งานระบบจองอุปกรณ์และสตูดิโออย่างเป็นลำดับ",
        "กำหนดแนวทางปฏิบัติงานสำหรับผู้ใช้แต่ละบทบาทให้ชัดเจน",
        "ใช้เป็นเอกสารอ้างอิงในการฝึกใช้งาน ส่งมอบระบบ และแก้ไขปัญหาเบื้องต้น",
    ])

    add_heading(doc, "ช.2 ขอบเขตการใช้งานระบบ", level=2)
    add_body(doc, "ระบบรองรับกระบวนการตั้งแต่การค้นหาทรัพยากร การสร้างรายการจอง การชำระเงินและแนบหลักฐาน การอนุมัติโดยเจ้าหน้าที่ การติดตามสถานะ และการปิดงาน")
    add_list(doc, [
        "แคตตาล็อกอุปกรณ์ สตูดิโอ แพ็กเกจ และบริการ",
        "ระบบตะกร้าและยืนยันการจองแบบหลายรายการ",
        "ระบบตรวจสอบความพร้อมใช้งานตามช่วงเวลา",
        "ระบบแนบสลิปและการตรวจสอบโดยเจ้าหน้าที่",
        "ระบบแจ้งเตือนอัตโนมัติและเอกสารใบเสนอราคา PDF",
    ])

    add_heading(doc, "ช.3 บทบาทผู้ใช้งาน", level=2)
    add_body(doc, "ลูกค้า: ค้นหาและจองทรัพยากร พร้อมติดตามสถานะของตนเอง")
    add_body(doc, "เจ้าหน้าที่: ตรวจสอบ อนุมัติ เปลี่ยนสถานะ และจัดการข้อมูลการปฏิบัติงาน")
    add_body(doc, "ผู้ดูแลระบบ: ดูแลระบบหลังบ้าน นำระบบขึ้นใช้งานจริง สำรองและกู้คืนข้อมูล")

    add_heading(doc, "ช.4 คู่มือการใช้งานสำหรับลูกค้า", level=2)
    add_heading(doc, "ช.4.1 การเข้าสู่ระบบ", level=3)
    add_list(doc, [
        "เปิดหน้าเว็บไซต์ระบบ",
        "เลือกเมนูเข้าสู่ระบบหรือสมัครสมาชิก",
        "กรอกอีเมลและรหัสผ่าน",
        "ตรวจสอบข้อมูลส่วนตัวให้ครบ โดยเฉพาะชื่อและเบอร์โทรศัพท์",
    ])
    add_image_placeholder(doc, 1, "หน้าเข้าสู่ระบบ/สมัครสมาชิก", "ท้ายหัวข้อ ช.4.1")

    add_heading(doc, "ช.4.2 การค้นหาและเลือกทรัพยากร", level=3)
    add_list(doc, [
        "เข้าเมนูแคตตาล็อก",
        "ใช้ช่องค้นหาและตัวกรองหมวดหมู่",
        "เปิดรายละเอียดเพื่อตรวจสอบราคา เงื่อนไข และความพร้อมใช้งาน",
        "กดเพิ่มลงตะกร้าสำหรับรายการที่ต้องการ",
    ])
    add_image_placeholder(doc, 2, "หน้าแคตตาล็อกและหน้ารายละเอียดรายการ", "ท้ายหัวข้อ ช.4.2")

    add_heading(doc, "ช.4.3 ขั้นตอนการจอง", level=3)
    add_list(doc, [
        "ตรวจสอบรายการในตะกร้า",
        "เลือกวันเริ่มต้นและวันสิ้นสุด",
        "กรอกข้อมูลติดต่อ/หมายเหตุ",
        "ตรวจสอบยอดรวมและยอดมัดจำ",
        "ยืนยันการจอง",
    ])
    add_image_placeholder(doc, 3, "หน้าตะกร้าและหน้ายืนยันการจอง", "ท้ายหัวข้อ ช.4.3")

    add_heading(doc, "ช.4.4 การชำระเงินและแนบสลิป", level=3)
    add_list(doc, [
        "เปิดเมนูรายการจองของฉัน (My Bookings)",
        "เลือกรายการที่ต้องการชำระเงิน",
        "โอนเงินตามข้อมูลบัญชีที่ระบบระบุ",
        "อัปโหลดไฟล์สลิปชำระเงิน",
        "รอเจ้าหน้าที่ตรวจสอบและยืนยัน",
    ])
    add_image_placeholder(doc, 4, "หน้า My Bookings และปุ่มอัปโหลดสลิป", "ท้ายหัวข้อ ช.4.4")

    add_heading(doc, "ช.4.5 การติดตามสถานะ", level=3)
    add_body(doc, "สถานะหลักในระบบ ได้แก่ Draft, Pending, Approved, Active, Overdue, Completed และ Cancelled")

    add_heading(doc, "ช.5 คู่มือการใช้งานสำหรับเจ้าหน้าที่", level=2)
    add_heading(doc, "ช.5.1 การตรวจสอบและอนุมัติรายการ", level=3)
    add_list(doc, [
        "เข้าสู่ระบบด้วยบัญชีเจ้าหน้าที่",
        "เปิด Staff Dashboard",
        "ตรวจสอบข้อมูลลูกค้า รายการทรัพยากร และช่วงเวลา",
        "อนุมัติหรือส่งกลับเพื่อแก้ไข",
    ])
    add_image_placeholder(doc, 5, "หน้า Staff Dashboard สำหรับตรวจสอบและอนุมัติ", "ท้ายหัวข้อ ช.5.1")

    add_heading(doc, "ช.5.2 การตรวจสลิปและยืนยันชำระเงิน", level=3)
    add_list(doc, [
        "เปิดรายการที่มีสลิปรอการตรวจสอบ",
        "เปรียบเทียบยอดโอนกับยอดที่ระบบกำหนด",
        "ยืนยันการชำระเงินเมื่อข้อมูลถูกต้อง",
    ])
    add_image_placeholder(doc, 6, "หน้าตรวจสลิปและปุ่มยืนยันการชำระเงิน", "ท้ายหัวข้อ ช.5.2")

    add_heading(doc, "ช.6 คู่มือสำหรับผู้ดูแลระบบ", level=2)
    add_heading(doc, "ช.6.1 คำสั่งตรวจระบบก่อนอัปเดต", level=3)
    add_code(doc, "python3 manage.py check")
    add_code(doc, "python3 manage.py test apps.store.tests_auth apps.store.tests_booking apps.store.tests_cancellation")

    add_heading(doc, "ช.6.2 การสำรองข้อมูล", level=3)
    add_code(doc, "./scripts/db_backup.sh")

    add_heading(doc, "ช.6.3 การทดสอบหลัง deploy", level=3)
    add_code(doc, "BASE_URL=https://mcotequipmentservices.mcot.net ./scripts/smoke_test.sh")

    add_heading(doc, "ช.6.4 การตรวจสุขภาพเซิร์ฟเวอร์", level=3)
    add_code(doc, "./scripts/health_check_vps.sh")

    add_heading(doc, "ช.6.5 การกู้คืนข้อมูล", level=3)
    add_code(doc, "CONFIRM_RESTORE=YES DRY_RUN=1 ./scripts/db_restore.sh backups/postgres/<file>.sql.gz")
    add_code(doc, "CONFIRM_RESTORE=YES ./scripts/db_restore.sh backups/postgres/<file>.sql.gz")

    add_heading(doc, "ช.7 ปัญหาที่พบบ่อยและแนวทางแก้ไข", level=2)
    add_body(doc, "เพิ่มสินค้าเข้าตะกร้าไม่ได้: ตรวจสอบสต็อกและช่วงเวลาใช้งาน")
    add_body(doc, "ยืนยันการจองไม่ได้: ตรวจสอบข้อมูลติดต่อและรายการที่ชนคิว")
    add_body(doc, "อัปโหลดสลิปแล้วสถานะไม่เปลี่ยน: รอเจ้าหน้าที่ตรวจสอบหรือประสานเจ้าหน้าที่")

    add_heading(doc, "ช.8 กฎการใช้งานและนโยบาย", level=2)
    add_list(doc, [
        "ผู้ใช้งานต้องกรอกข้อมูลให้ถูกต้องก่อนยืนยันรายการ",
        "การอนุมัติขึ้นอยู่กับการตรวจสอบของเจ้าหน้าที่",
        "การคืนล่าช้าอาจมีค่าปรับตามเงื่อนไข",
        "เมนูบางส่วนสงวนสิทธิ์สำหรับเจ้าหน้าที่หรือผู้ดูแลระบบ",
    ])

    add_heading(doc, "ช.9 ช่องทางติดต่อ", level=2)
    add_body(doc, "ทีม IT Support: it-support@rental.mcot.net")
    add_body(doc, "โทรภายใน: 6000 ต่อ 123")
    add_body(doc, "ผู้ประสานงานหน่วยงาน: ให้ระบุข้อมูลจริงก่อนส่งมอบ")

    add_heading(doc, "ช.10 รายการภาพประกอบที่ต้องใส่", level=2)
    add_list(doc, [
        "ภาพ ช.1 หน้าเข้าสู่ระบบ/สมัครสมาชิก",
        "ภาพ ช.2 หน้าแคตตาล็อกและรายละเอียดรายการ",
        "ภาพ ช.3 หน้าตะกร้าและยืนยันการจอง",
        "ภาพ ช.4 หน้า My Bookings และอัปโหลดสลิป",
        "ภาพ ช.5 หน้า Staff Dashboard อนุมัติรายการ",
        "ภาพ ช.6 หน้าตรวจสลิปและยืนยันชำระเงิน",
    ])

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
