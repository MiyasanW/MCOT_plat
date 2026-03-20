from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "PROJECT_chapter3_methodology.docx"


def set_a4_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)


def set_th_sarabun(run, size=14, bold=False):
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(size)
    run.bold = bold


def add_center_title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_th_sarabun(run, size=size, bold=True)


def add_heading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_th_sarabun(run, size=16 if level == 2 else 14, bold=True)


def add_body(doc, text, first_line_indent_cm=0.75):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_th_sarabun(run, size=14)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            set_th_sarabun(run, size=14)


def add_table_with_headers(doc, title, headers, rows):
    add_heading(doc, title, level=3)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                set_th_sarabun(r, size=13, bold=True)

    for row_data in rows:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
            for p in row[i].paragraphs:
                for r in p.runs:
                    set_th_sarabun(r, size=12)


def main():
    doc = Document()
    set_a4_margins(doc)

    add_center_title(doc, "บทที่ 3")
    add_center_title(doc, "วิธีการดำเนินงาน")
    doc.add_paragraph()

    add_heading(doc, "3.1 ระเบียบวิธีที่ใช้ในการพัฒนา", level=2)
    add_body(
        doc,
        "โครงงานนี้ใช้แนวทางการพัฒนาแบบเป็นรอบและเพิ่มความสมบูรณ์อย่างต่อเนื่อง "
        "(Iterative and Incremental Development) โดยเริ่มจากการวิเคราะห์ปัญหาจริงของหน่วยงาน กำหนดขอบเขต "
        "ที่จำเป็น พัฒนาโมดูลหลักก่อน แล้วรับข้อเสนอแนะจากผู้ใช้งานจริงเพื่อนำไปปรับปรุงในรอบถัดไป "
        "แนวทางดังกล่าวเหมาะกับระบบจองที่มีเงื่อนไขทางธุรกิจเปลี่ยนได้ระหว่างดำเนินโครงการ"
    )

    add_body(doc, "ลำดับการพัฒนาหลักของโครงงานมีดังนี้", first_line_indent_cm=0)
    add_bullets(
        doc,
        [
            "วิเคราะห์ความต้องการของระบบและจัดลำดับความสำคัญของฟังก์ชัน",
            "ออกแบบข้อมูลและกระบวนการทำงานให้สอดคล้องกับงานจริง",
            "พัฒนาโมดูลหลักทีละส่วนและทดสอบร่วมกับผู้ใช้งาน",
            "ปรับปรุงความถูกต้อง เสถียรภาพ และประสบการณ์ผู้ใช้",
            "นำระบบขึ้นใช้งานจริงและตรวจสอบผลหลัง deploy",
        ],
    )

    add_heading(doc, "3.2 การเก็บรวบรวมความต้องการ", level=2)
    add_body(
        doc,
        "การเก็บรวบรวมความต้องการดำเนินการจากหลายช่องทางเพื่อให้ครอบคลุมทั้งฝั่งผู้ใช้งานและฝั่งปฏิบัติการ "
        "ประกอบด้วยการสัมภาษณ์เจ้าหน้าที่ การสังเกตขั้นตอนงานเดิม การทบทวนเอกสาร และการทดลองใช้งานต้นแบบ "
        "จากนั้นจึงสรุปเป็นความต้องการเชิงหน้าที่และเชิงคุณภาพ"
    )

    add_heading(doc, "3.2.1 ความต้องการเชิงหน้าที่ (Functional Requirements)", level=3)
    add_body(
        doc,
        "ความต้องการเชิงหน้าที่ของระบบถูกกำหนดจากงานจริงของผู้ใช้งาน โดยแต่ละข้อระบุอินพุต เงื่อนไข และผลลัพธ์ที่ต้องเกิดขึ้นอย่างชัดเจน"
    )
    add_table_with_headers(
        doc,
        "ตาราง 3.1 สรุปความต้องการเชิงหน้าที่",
        ["รหัส", "ความต้องการ", "อินพุต/เงื่อนไข", "ผลลัพธ์"],
        [
            ["FR-01", "ค้นหาและกรองทรัพยากร", "คำค้นหา/หมวดหมู่/ช่วงเวลา", "แสดงรายการที่ตรงเงื่อนไข"],
            ["FR-02", "ตรวจสอบ availability", "id รายการ วันเริ่ม-สิ้นสุด จำนวน", "แจ้งพร้อมใช้/ไม่พร้อมใช้"],
            ["FR-03", "สร้างรายการจองจากตะกร้า", "ข้อมูลตะกร้าและผู้จอง", "สร้าง Booking/BookingItem"],
            ["FR-04", "คำนวณราคาและมัดจำ", "ราคา วันเช่า โปรโมชัน %มัดจำ", "ยอดรวม ส่วนลด มัดจำ"],
            ["FR-05", "แนบสลิปและตรวจสอบ", "ไฟล์สลิป + รายการจอง", "อัปเดตสถานะและแจ้งเตือน"],
            ["FR-06", "จัดการสถานะตาม workflow", "สถานะเดิม + action + สิทธิ์", "เปลี่ยนสถานะพร้อมบันทึกประวัติ"],
            ["FR-07", "ส่งอีเมลแจ้งเตือน", "เหตุการณ์สำคัญของระบบ", "ส่งอีเมลให้ผู้เกี่ยวข้อง"],
            ["FR-08", "สร้างใบเสนอราคา PDF", "ข้อมูลจองและยอดเงิน", "ได้ไฟล์ PDF สำหรับแนบ/ดาวน์โหลด"],
        ],
    )

    doc.add_paragraph()
    add_heading(doc, "3.2.2 ความต้องการเชิงคุณภาพ (Non-functional Requirements)", level=3)
    add_body(
        doc,
        "ความต้องการเชิงคุณภาพถูกกำหนดเพื่อให้ระบบมีเสถียรภาพ ปลอดภัย และดูแลรักษาได้ในระยะยาว โดยเน้นตัวชี้วัดที่ตรวจสอบได้"
    )
    add_table_with_headers(
        doc,
        "ตาราง 3.2 สรุปความต้องการเชิงคุณภาพ",
        ["รหัส", "ด้านคุณภาพ", "เกณฑ์/ตัวชี้วัด", "แนวทางที่ใช้"],
        [
            ["NFR-01", "Data Integrity", "ไม่เกิดคิวจองซ้อน", "ตรวจ availability + server-side validation"],
            ["NFR-02", "Security", "HTTPS + CSRF + RBAC", "ใช้ secure settings และตรวจสิทธิ์ทุกจุดสำคัญ"],
            ["NFR-03", "Availability", "บริการหลักทำงานต่อเนื่อง", "VPS + service monitoring + health check"],
            ["NFR-04", "Performance", "ตอบสนองเร็วในหน้าหลัก", "ลด query ไม่จำเป็น แยก logic ใน service"],
            ["NFR-05", "Maintainability", "แก้ไขง่าย กระทบต่ำ", "แยกโค้ดเป็น views/services/models"],
            ["NFR-06", "Auditability", "ตรวจย้อนหลังการเปลี่ยนสถานะได้", "บันทึก history และ logs"],
        ],
    )

    add_heading(doc, "3.4 การออกแบบสถาปัตยกรรมระบบ", level=2)
    add_body(
        doc,
        "สถาปัตยกรรมระบบถูกออกแบบแบบแยกชั้น (Layered Architecture) เพื่อแยกความรับผิดชอบของแต่ละส่วนให้ชัดเจน "
        "ลดการพึ่งพากันของโค้ด และทำให้การปรับปรุงในอนาคตทำได้เป็นจุด ๆ โดยไม่กระทบทั้งระบบ"
    )

    add_body(
        doc,
        "ในเชิงการไหลของคำขอ (request flow) ผู้ใช้จะเริ่มจากหน้าเว็บ (Template) ส่งคำขอไปยัง View จากนั้น View จะเรียก "
        "Service เพื่อประมวลผลตรรกะธุรกิจ เช่น ตรวจ availability คำนวณราคา/มัดจำ และตรวจสิทธิ์ ก่อนบันทึกหรืออ่านข้อมูลผ่าน Model "
        "และส่งผลกลับมายัง Template เพื่อแสดงผล"
    )

    add_heading(doc, "3.4.1 โครงสร้างเชิงชั้นของระบบ", level=3)
    add_bullets(
        doc,
        [
            "Presentation Layer: รับผิดชอบส่วนติดต่อผู้ใช้ เช่น หน้า catalog, cart, my bookings และ staff dashboard",
            "Application Layer (Views/URLs): รับ request, ตรวจสิทธิ์, จัดการลำดับ flow และเลือก service ที่เกี่ยวข้อง",
            "Service Layer: รวมตรรกะธุรกิจที่ใช้ซ้ำ เช่น PricingService, BookingService, NotificationService, AvailabilityService",
            "Data Layer (Models/DB): จัดเก็บข้อมูลหลักของระบบและบังคับกฎระดับข้อมูล เช่น ความสัมพันธ์และ validation",
        ],
    )

    add_heading(doc, "3.4.2 เหตุผลที่เลือกสถาปัตยกรรมนี้", level=3)
    add_bullets(
        doc,
        [
            "รองรับการเปลี่ยน requirement ระหว่างพัฒนาได้ง่าย เพราะตรรกะหลักอยู่ใน service layer",
            "เพิ่มความสามารถในการทดสอบ (testability) โดยทดสอบ business logic แยกจากหน้าเว็บได้",
            "ลดความเสี่ยง regression เพราะการแก้ไขส่วนแสดงผลไม่จำเป็นต้องแก้ตรรกะข้อมูล",
            "ทำให้การส่งมอบงานต่อทีมอื่นในอนาคตทำได้ง่ายขึ้นจากโครงสร้างที่อ่านเข้าใจได้เร็ว",
        ],
    )

    add_heading(doc, "3.4.3 การแม็ปสถาปัตยกรรมกับโมดูลจริง", level=3)
    add_bullets(
        doc,
        [
            "ชั้นแสดงผล: templates/booking, templates/staff, templates/store",
            "ชั้นควบคุม: apps/store/views/booking.py, staff.py, notification.py",
            "ชั้นบริการ: apps/store/services/pricing_service.py, booking_service.py, notification_service.py",
            "ชั้นข้อมูล: apps/store/models.py และฐานข้อมูล PostgreSQL/SQLite ตามสภาพแวดล้อม",
        ],
    )

    add_heading(doc, "3.5 การออกแบบฐานข้อมูล", level=2)
    add_body(
        doc,
        "ฐานข้อมูลถูกออกแบบให้รองรับทั้งข้อมูลตั้งต้น (master data) และข้อมูลธุรกรรม (transaction data) โดยเน้น "
        "ความถูกต้องของความสัมพันธ์ระหว่างตารางและความสามารถในการตรวจสอบย้อนหลังของรายการจอง"
    )

    add_heading(doc, "3.5.1 โครงสร้างข้อมูลหลัก", level=3)
    add_body(
        doc,
        "กลุ่มข้อมูลทรัพยากรประกอบด้วยข้อมูลอุปกรณ์ สตูดิโอ แพ็กเกจ และบริการ ซึ่งเป็นต้นทางของการสร้างรายการจอง "
        "แต่ละรายการมีคุณสมบัติ ราคา สถานะพร้อมใช้งาน และเงื่อนไขการใช้งานที่แตกต่างกัน"
    )
    add_body(
        doc,
        "กลุ่มข้อมูลธุรกรรมประกอบด้วย Booking และ BookingItem โดย Booking ทำหน้าที่เป็นหัวเอกสาร "
        "ส่วน BookingItem เป็นรายละเอียดรายการย่อย ทำให้รองรับการจองหลายรายการในหนึ่งธุรกรรม และสามารถคำนวณราคาต่อรายการได้"
    )
    add_body(
        doc,
        "กลุ่มข้อมูลผู้ใช้งานประกอบด้วย User/Profile/Staff เพื่อรองรับการกำหนดบทบาทและสิทธิ์การทำงาน "
        "เช่น ลูกค้าสร้างรายการได้ เจ้าหน้าที่อนุมัติและเปลี่ยนสถานะได้"
    )

    add_heading(doc, "3.5.2 ความสัมพันธ์ของข้อมูล (Data Relationship)", level=3)
    add_bullets(
        doc,
        [
            "User 1 คน สามารถมี Booking ได้หลายรายการ (One-to-Many)",
            "Booking 1 รายการ มี BookingItem ได้หลายรายการ (One-to-Many)",
            "BookingItem แต่ละรายการอ้างอิงทรัพยากรที่ถูกจอง เช่น Product/Studio/Package",
            "Notification เชื่อมกับ Booking เพื่อแจ้งเหตุการณ์สำคัญตามสถานะงาน",
        ],
    )

    add_heading(doc, "3.5.3 กฎความถูกต้องของข้อมูล", level=3)
    add_bullets(
        doc,
        [
            "ห้ามสร้าง booking หากรายการในตะกร้าไม่มีอยู่จริงหรือไม่พร้อมใช้งานในช่วงเวลาที่เลือก",
            "เปอร์เซ็นต์มัดจำต้องอยู่ในช่วงที่กำหนดและใช้คำนวณยอดมัดจำได้อย่างสอดคล้อง",
            "การเปลี่ยนสถานะต้องเป็นไปตาม workflow ที่กำหนด เช่น pending -> approved -> paid -> completed",
            "ข้อมูลสลิปชำระเงินต้องเชื่อมกับ booking ที่ถูกต้องและตรวจสอบย้อนหลังได้",
        ],
    )

    add_heading(doc, "3.5.4 แนวทางรองรับการขยายระบบ", level=3)
    add_bullets(
        doc,
        [
            "รองรับการเพิ่มประเภททรัพยากรใหม่โดยไม่กระทบโครงสร้างธุรกรรมหลัก",
            "รองรับการเพิ่มฟิลด์เชิงวิเคราะห์ในอนาคต เช่น ช่องทางการจองหรือกลุ่มลูกค้า",
            "รองรับการเชื่อมต่อกับระบบภายนอกในระยะถัดไปผ่าน service layer และ key mappings",
        ],
    )

    add_heading(doc, "3.6 เครื่องมือและภาษาที่ใช้ในการพัฒนา", level=2)
    add_body(
        doc,
        "เพื่อให้การพัฒนาระบบเป็นไปอย่างมีประสิทธิภาพและสามารถดูแลรักษาต่อได้ในสภาพแวดล้อมจริง โครงงานนี้ได้เลือกใช้ "
        "เครื่องมือและเทคโนโลยีที่เหมาะสมกับลักษณะงานจองทรัพยากรขององค์กร"
    )

    add_heading(doc, "3.6.1 ภาษาที่ใช้ในการพัฒนา", level=3)
    add_bullets(
        doc,
        [
            "Python: ใช้พัฒนาระบบฝั่งหลังบ้าน (Backend) และตรรกะธุรกิจหลัก",
            "HTML/CSS/JavaScript: ใช้พัฒนาส่วนติดต่อผู้ใช้ (Frontend) และพฤติกรรมหน้าเว็บ",
            "SQL: ใช้จัดการข้อมูลเชิงสัมพันธ์และคำสั่งตรวจสอบข้อมูลในฐานข้อมูล",
        ],
    )

    add_heading(doc, "3.6.2 Framework และไลบรารีสำคัญ", level=3)
    add_bullets(
        doc,
        [
            "Django 4.2: เฟรมเวิร์กหลักสำหรับพัฒนาเว็บแอปพลิเคชัน",
            "Django Simple History: ใช้บันทึกประวัติการเปลี่ยนแปลงข้อมูลสำคัญ",
            "ReportLab และ PyPDF2: ใช้สร้างและจัดการเอกสาร PDF เช่น ใบเสนอราคา",
            "python-docx: ใช้สร้างเอกสาร Word สำหรับรายงานและบทโครงงาน",
        ],
    )

    add_heading(doc, "3.6.3 เครื่องมือพัฒนาและปฏิบัติการ", level=3)
    add_bullets(
        doc,
        [
            "Visual Studio Code: เครื่องมือหลักสำหรับพัฒนาและแก้ไขโค้ด",
            "Git และ GitHub: ใช้ควบคุมเวอร์ชันและติดตามประวัติการเปลี่ยนแปลง",
            "SQLite: ใช้ในสภาพแวดล้อมพัฒนา (development)",
            "PostgreSQL: ใช้ในสภาพแวดล้อม production",
            "Gunicorn + Nginx บน VPS: ใช้ให้บริการระบบจริง",
            "Shell Scripts (เช่น smoke test, health check, db backup/restore): ใช้ช่วยงานดูแลระบบ",
        ],
    )

    add_heading(doc, "3.7 การพัฒนาและการทดสอบ", level=2)
    add_body(
        doc,
        "การพัฒนาดำเนินแบบเป็นรอบ โดยแต่ละรอบจะมีการพัฒนา ทดสอบ และรับข้อเสนอแนะก่อนส่งต่อรอบถัดไป "
        "ทำให้แก้ไขปัญหาได้เร็วและลดผลกระทบสะสม"
    )

    add_body(doc, "แนวทางการทดสอบที่ใช้ในโครงงาน", first_line_indent_cm=0)
    add_bullets(
        doc,
        [
            "Unit Test: ทดสอบตรรกะย่อย เช่น คำนวณราคา มัดจำ และเงื่อนไขสถานะ",
            "Integration Test: ทดสอบการทำงานร่วมกันระหว่างโมดูล",
            "Functional Test: ทดสอบตามขั้นตอนใช้งานจริงของลูกค้าและเจ้าหน้าที่",
            "Smoke Test: ทดสอบภาพรวมหลัง deploy บน production",
        ],
    )

    add_heading(doc, "3.7.1 กรณีทดสอบสำคัญและผลที่คาดหวัง", level=3)
    add_body(
        doc,
        "เพื่อให้มั่นใจว่าระบบทำงานได้ตามข้อกำหนด มีการกำหนดกรณีทดสอบสำคัญครอบคลุมเส้นทางใช้งานหลัก โดยสรุปได้ดังนี้"
    )
    add_table_with_headers(
        doc,
        "ตาราง 3.3 ตัวอย่างกรณีทดสอบสำคัญ",
        ["รหัสทดสอบ", "กรณีทดสอบ", "ผลที่คาดหวัง"],
        [
            ["TC-01", "สร้าง booking จาก cart ปกติ", "สร้างรายการสำเร็จและคำนวณยอดถูกต้อง"],
            ["TC-02", "id ใน cart ไม่ถูกต้อง", "ปฏิเสธรายการพร้อมข้อความอธิบาย"],
            ["TC-03", "จองช่วงเวลาชนกัน", "แจ้งไม่พร้อมใช้งานและไม่ยืนยันรายการ"],
            ["TC-04", "ผู้ใช้ทั่วไปเรียก staff action", "ระบบปฏิเสธสิทธิ์"],
            ["TC-05", "แนบสลิปและยืนยันชำระเงิน", "บันทึกหลักฐานและเปลี่ยนสถานะตาม workflow"],
        ],
    )

    add_heading(doc, "3.8 การนำระบบขึ้นใช้งานจริง", level=2)
    add_body(
        doc,
        "การนำระบบขึ้นใช้งานจริง (Production Deployment) ถูกออกแบบให้เป็นกระบวนการที่ตรวจสอบย้อนกลับได้ "
        "ลดความเสี่ยงจาก human error และมีจุดควบคุมก่อน-หลัง deploy อย่างชัดเจน"
    )

    add_heading(doc, "3.8.1 การเตรียมความพร้อมก่อน deploy", level=3)
    add_bullets(
        doc,
        [
            "ยืนยัน branch/commit ที่จะนำขึ้นใช้งาน และตรวจความถูกต้องของโค้ดล่าสุด",
            "ตรวจไฟล์ตั้งค่า environment variables ให้ครบตาม production",
            "รันคำสั่งตรวจระบบและชุดทดสอบที่เกี่ยวข้องกับฟีเจอร์ที่แก้ไข",
            "ตรวจรายการเปลี่ยนแปลงฐานข้อมูลและแผน migration",
            "สำรองฐานข้อมูลก่อนเริ่ม deploy ทุกครั้ง",
        ],
    )

    add_heading(doc, "3.8.2 ขั้นตอนการ deploy บน VPS", level=3)
    add_bullets(
        doc,
        [
            "ดึงโค้ดล่าสุดบนเครื่อง VPS และติดตั้ง dependency ให้ตรงกับ requirements",
            "รัน migration เพื่อปรับ schema ฐานข้อมูลให้สอดคล้องกับโค้ดเวอร์ชันใหม่",
            "รัน collectstatic เพื่ออัปเดตไฟล์ static ที่หน้าเว็บใช้งาน",
            "restart/reload บริการแอปพลิเคชันและเว็บเซิร์ฟเวอร์ เช่น gunicorn และ nginx",
            "ตรวจสถานะบริการว่าทำงานครบและไม่มี service ใดล้มเหลว",
        ],
    )

    add_heading(doc, "3.8.3 การทดสอบหลัง deploy (Post-deployment Verification)", level=3)
    add_bullets(
        doc,
        [
            "ทดสอบหน้าแรก หน้าล็อกอิน หน้าแคตตาล็อก และเส้นทางจองหลักแบบ end-to-end",
            "ทดสอบ flow สำคัญ: สร้าง booking, แนบสลิป, เปลี่ยนสถานะ, ส่งอีเมลแจ้งเตือน",
            "ตรวจ log ของแอปและเว็บเซิร์ฟเวอร์เพื่อยืนยันว่าไม่มี error ผิดปกติ",
            "รัน smoke test และ health check script เพื่อยืนยันความพร้อมให้บริการ",
        ],
    )

    add_heading(doc, "3.8.4 แผนรับมือเมื่อ deploy ไม่สำเร็จ", level=3)
    add_bullets(
        doc,
        [
            "หาก migration ล้มเหลว ให้หยุดการเปลี่ยนแปลงและย้อนกลับตามขั้นตอนที่กำหนด",
            "หากบริการไม่ขึ้น ให้ตรวจ config, dependency, และ log ก่อนตัดสินใจ rollback",
            "หากเกิดผลกระทบผู้ใช้ ให้คืนค่าจาก backup ที่ทำไว้ก่อน deploy",
            "บันทึก incident และสรุปบทเรียนเพื่อป้องกันปัญหาซ้ำใน release ถัดไป",
        ],
    )

    add_heading(doc, "3.9 สรุปวิธีการดำเนินงาน", level=2)
    add_body(
        doc,
        "วิธีการดำเนินงานของโครงงานนี้เน้นการพัฒนาแบบมีรอบตรวจสอบจริง ควบคู่กับการแยกสถาปัตยกรรมให้ดูแลรักษาง่าย "
        "และการทดสอบหลายระดับ เพื่อให้ระบบที่พัฒนาสามารถใช้งานได้จริงในองค์กร มีความปลอดภัย และพร้อมต่อยอดในระยะถัดไป"
    )

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
