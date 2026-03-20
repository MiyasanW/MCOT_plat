from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "PROJECT_chapter4_workflow.docx"


def set_cell_background(cell, fill_color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def set_a4_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'TH Sarabun New'


def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.75)
    for run in p.runs:
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(14)


def main():
    doc = Document()
    set_a4_margins(doc)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("บทที่ 4: ผลการปฏิบัติงาน")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = 'TH Sarabun New'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("การไหลงานและสถานะของระบบจองอุปกรณ์")
    run.font.size = Pt(14)
    run.font.name = 'TH Sarabun New'

    doc.add_paragraph()

    # Workflow Description
    add_heading(doc, "การไหลงานหลักของระบบ", level=2)
    add_paragraph(
        doc,
        "ระบบจองอุปกรณ์ถ่ายทำและสตูดิโอมีการไหลงานหลักประกอบด้วย 4 สถานการณ์ หรือ "
        "workflow ที่แตกต่างกัน โดยแต่ละสถานการณ์มีลำดับของสถานะ (status) และการแจ้งเตือน "
        "ที่เกี่ยวข้อง ดังรายละเอียดต่อไปนี้"
    )

    doc.add_paragraph()

    # Workflow Table
    add_heading(doc, "ตาราง 4.1 การไหลงานของระบบตั้งแต่ต้นจนท้าย", level=3)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'

    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = "สถานการณ์"
    header_cells[1].text = "ผู้ดำเนินการ"
    header_cells[2].text = "สถานะเดิม → สถานะใหม่"
    header_cells[3].text = "การแจ้งเตือน"

    for cell in header_cells:
        set_cell_background(cell, "4472C4")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'TH Sarabun New'
                run.font.size = Pt(12)

    # Workflow rows
    workflows = [
        [
            "1. ลูกค้าค้นหาและเลือก\nจองทรัพยากร",
            "ลูกค้า",
            "Draft / Pending → Pending",
            "อีเมลยืนยันการจอง\n+ ใบเสนอราคา PDF",
        ],
        [
            "2. เจ้าหน้าที่ตรวจสอบ\nและอนุมัติ",
            "เจ้าหน้าที่",
            "Pending → Approved",
            "อีเมลแจ้งอนุมัติ\nให้ลูกค้า",
        ],
        [
            "3. ลูกค้าชำระเงิน\nและอัปโหลดสลิป",
            "ลูกค้า / เจ้าหน้าที่",
            "Approved → Paid",
            "อีเมลยืนยันชำระเงิน\n+ รายละเอียดการนำส่ง",
        ],
        [
            "4. จัดเตรียม นำส่ง\nและปิดงาน",
            "เจ้าหน้าที่",
            "Paid → Completed",
            "อีเมลยืนยันการนำส่ง\nสิ้นสุดการจอง",
        ],
    ]

    for row_idx, (scenario, actor, status, notification) in enumerate(workflows):
        row = table.add_row()
        row.cells[0].text = scenario
        row.cells[1].text = actor
        row.cells[2].text = status
        row.cells[3].text = notification

        # Handle alternating colors
        for i, cell in enumerate(row.cells):
            if row_idx % 2 == 0:
                set_cell_background(cell, "D9E8F5")

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'TH Sarabun New'
                    run.font.size = Pt(11)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    # State Transition Details
    add_heading(doc, "สถานะและการเปลี่ยนผ่านของการจอง", level=2)

    state_table = doc.add_table(rows=1, cols=3)
    state_table.style = 'Light Grid Accent 1'

    # Header
    state_header = state_table.rows[0].cells
    state_header[0].text = "สถานะ\n(Status)"
    state_header[1].text = "หมายความว่า"
    state_header[2].text = "การดำเนินการต่อไป"

    for cell in state_header:
        set_cell_background(cell, "70AD47")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'TH Sarabun New'
                run.font.size = Pt(12)

    states_data = [
        [
            "Draft",
            "ร่างการจอง: ลูกค้าเพิ่มรายการลงตะกร้า\nแต่ยังไม่ยืนยัน",
            "ลูกค้าสามารถแก้ไขตะกร้า\nหรือยืนยันการจอง",
        ],
        [
            "Pending",
            "รอการตรวจสอบ: เจ้าหน้าที่ยังไม่ได้\nตรวจสอบและอนุมัติ",
            "เจ้าหน้าที่ตรวจสอบและ\nตัดสินใจอนุมัติ",
        ],
        [
            "Approved",
            "ได้รับการอนุมัติ: เจ้าหน้าที่อนุมัติแล้ว\nรอการชำระเงิน",
            "ลูกค้าชำระเงินมัดจำ\nและอัปโหลดสลิป",
        ],
        [
            "Paid",
            "ชำระเงินแล้ว: เจ้าหน้าที่ยืนยินชำระ\nพร้อมจัดเตรียมอุปกรณ์",
            "เจ้าหน้าที่เตรียม อุปกรณ์\nและนำส่งตามวันธรรม",
        ],
        [
            "Active",
            "ใช้งานอยู่: ลูกค้าได้รับอุปกรณ์\nและอยู่ในช่วงเช่า",
            "หลังสิ้นสุดวันเช่า\nลูกค้าคืนอุปกรณ์",
        ],
        [
            "Completed",
            "สิ้นสุด: ลูกค้าคืนอุปกรณ์และ\nเจ้าหน้าที่ตรวจสภาพ",
            "ปิดการจองทั้งหมด\nเก็บบันทึกประวัติ",
        ],
        [
            "Cancelled",
            "ยกเลิก: ลูกค้าหรือเจ้าหน้าที่\nยกเลิกการจองใน course",
            "คืนเงินมัดจำ (หากเหมาะสม)\nบันทึกสาเหตุการยกเลิก",
        ],
    ]

    for row_idx, (status, meaning, action) in enumerate(states_data):
        row = state_table.add_row()
        row.cells[0].text = status
        row.cells[1].text = meaning
        row.cells[2].text = action

        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, "E2EFD9")

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'TH Sarabun New'
                    run.font.size = Pt(11)

    doc.add_paragraph()

    # Email Notifications Table
    add_heading(doc, "ตาราง 4.3 การส่งอีเมลแจ้งเตือนอัตโนมัติ", level=2)

    email_table = doc.add_table(rows=1, cols=4)
    email_table.style = 'Light Grid Accent 1'

    email_header = email_table.rows[0].cells
    email_header[0].text = "เหตุการณ์"
    email_header[1].text = "เมื่อไร"
    email_header[2].text = "ส่งไปที่"
    email_header[3].text = "เนื้อหาหลัก"

    for cell in email_header:
        set_cell_background(cell, "C55A11")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'TH Sarabun New'

    email_data = [
        [
            "สร้างการจอง",
            "เมื่อลูกค้าส่ง\nการจอง",
            "ลูกค้า",
            "ใบเสนอราคา (PDF)\nรายละเอียดการจอง\nบัญชีสำหรับชำระเงิน",
        ],
        [
            "อนุมัติการจอง",
            "เมื่อเจ้าหน้าที่\nอนุมัติ",
            "ลูกค้า",
            "การจองได้รับการยืนยัน\nยอดมัดจำที่ต้องชำระ\nวิธีการชำระเงิน",
        ],
        [
            "อัปโหลดสลิป",
            "เมื่อลูกค้า\nอัปโหลดสลิป",
            "เจ้าหน้าที่",
            "แจ้งว่ามีสลิปรอ\nการตรวจสอบ\nลิงก์ไปหน้าอนุมัติ",
        ],
        [
            "ยืนยันชำระเงิน",
            "เมื่อเจ้าหน้าที่\nยืนยินชำระ",
            "ลูกค้า",
            "การชำระเงินได้รับการยืนยัน\nรายละเอียดการนำส่ง\nเวลาและสถานที่รับ",
        ],
        [
            "ปิดการจอง",
            "เมื่อลูกค้าคืน\nและเรียบร้อย",
            "ลูกค้า",
            "บันทึกการจอง\nใบเสร็จครบถ้วน\nขอบคุณและข้อเสนอแนะ",
        ],
    ]

    for row_idx, (event, when, recipient, content) in enumerate(email_data):
        row = email_table.add_row()
        row.cells[0].text = event
        row.cells[1].text = when
        row.cells[2].text = recipient
        row.cells[3].text = content

        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, "FCE4D6")

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'TH Sarabun New'
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # Summary
    add_heading(doc, "สรุป", level=2)
    add_paragraph(
        doc,
        "ระบบจองอุปกรณ์ได้ถูกออกแบบให้มีการไหลงานที่ชัดเจนและเป็นระเบียบ "
        "โดยมีการแจ้งเตือนอัตโนมัติผ่านอีเมลในแต่ละขั้นตอนสำคัญ "
        "เพื่อให้ลูกค้าและเจ้าหน้าที่ได้รับข้อมูลทันท่วงที และสามารถติดตามสถานะการจอง "
        "ได้อย่างต่อเนื่อง ซึ่งช่วยลดความเข้าใจผิดและเพิ่มประสิทธิภาพของการบริหารจัดการอุปกรณ์"
    )

    doc.save(OUTPUT)
    print(f"✅ สร้าง {OUTPUT} สำเร็จ")


if __name__ == "__main__":
    main()
